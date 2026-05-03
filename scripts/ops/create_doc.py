from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    if level == 1 and heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(46, 80, 144)
    return heading


def shade_cell(cell, color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_monospace_block(doc, text, size=9):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        shade_cell(table.rows[0].cells[idx], "D5E8F0")
    for row_idx, row in enumerate(rows, 1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value
    return table


# This helper generates a lightweight orientation document. Keep it aligned
# with the current self-scheduling runtime, not the retired walk scheduler.
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Project Architecture Overview")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(46, 80, 144)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("NYC Air Quality Field Monitoring System")
run.font.size = Pt(16)
run.font.italic = True

doc.add_page_break()

add_heading(doc, "Table of Contents", level=1)
add_bullets(doc, [
    "Executive Summary",
    "System Overview",
    "System Architecture",
    "Core Runtime Components",
    "Data Flow",
    "Key Data Sources",
    "Technology Stack",
    "Key Interconnections",
    "Quick Reference",
    "For New Team Members",
])

doc.add_page_break()

add_heading(doc, "Executive Summary", level=1)
doc.add_paragraph(
    "A real-time field monitoring platform for air quality research where collectors "
    "walk predetermined routes with GPS trackers. The current system supports direct "
    "self-scheduling, weather advisory refreshes, Drive sync, and real-time dashboard "
    "monitoring. The legacy algorithmic walk scheduler is retired from the active "
    "runtime path."
)

add_heading(doc, "System Overview", level=1)
add_bullets(doc, [
    "What: Real-time field monitoring for air quality research",
    "Who: NYC field collectors walking predetermined routes",
    "How: Collector self-scheduling with weather advisories and server-side rebuilds",
    "Where: Deployed on Google Cloud Run with GCS-backed persisted state",
])

add_heading(doc, "System Architecture", level=1)
add_monospace_block(doc, """FIELD LAYER
- Collector phones / GPS tracks
- Google Drive walk data uploads
- Google Sheets forecast updates
       |
       v
[serve.py]  central HTTP/API runtime
       |-- build_weather.py -> weather.json
       |-- build_dashboard.py -> dashboard.html
       `-- schedule APIs -> schedule_output.json""", size=9)

add_heading(doc, "Core Runtime Components", level=1)

add_heading(doc, "1. serve.py - HTTP Server and Data Hub", level=2)
doc.add_paragraph().add_run("Purpose: ").bold = True
doc.paragraphs[-1].add_run("Central coordinator for dashboard serving, schedule APIs, Drive sync, rebuild triggers, and upload handling.")
add_bullets(doc, [
    "Serves dashboard files and JSON API responses",
    "Reads and writes schedule_output.json through self-scheduling endpoints",
    "Triggers weather and dashboard rebuilds",
    "Restores and persists runtime state through GCS",
    "Handles Drive polling or GAS-triggered Drive sync",
])

add_heading(doc, "2. Self-Scheduling Runtime", level=2)
doc.add_paragraph().add_run("Purpose: ").bold = True
doc.paragraphs[-1].add_run("Lets collectors claim and update schedule slots directly while preserving basic conflict rules.")
add_bullets(doc, [
    "Claims one backpack/date/time-of-day slot at a time",
    "Blocks duplicate backpack slots for the same date and time of day",
    "Blocks collector double-booking in the same date and time-of-day slot",
    "Stores assignments in schedule_output.json",
    "Uses weather as advisory data only",
    "Retired POST /api/rerun, /api/rerun/a, and /api/rerun/b return 410 Gone",
])

add_heading(doc, "3. build_weather.py - Weather Advisory Builder", level=2)
doc.add_paragraph().add_run("Purpose: ").bold = True
doc.paragraphs[-1].add_run("Builds weather.json from forecast source data for dashboard advisory display.")
add_bullets(doc, [
    "Reads forecast tabs from Google Sheets",
    "Normalizes date/time-of-day advisory data",
    "Writes data/outputs/site/weather.json",
])

add_heading(doc, "4. build_dashboard.py - Dashboard Generator", level=2)
doc.add_paragraph().add_run("Purpose: ").bold = True
doc.paragraphs[-1].add_run("Creates the interactive HTML dashboard for monitoring, scheduling, uploads, and weather visibility.")
add_bullets(doc, [
    "Embeds route geometry and walk log data",
    "Embeds self-scheduling and weather data snapshots",
    "Builds the Leaflet.js map and dashboard UI",
    "Writes data/outputs/site/dashboard.html",
])

add_heading(doc, "5. Retired Scheduler Artifacts", level=2)
doc.add_paragraph(
    "The legacy walk scheduler, map builder, and transit matrix scripts live under "
    "pipelines/_retired/. They remain historical/reference tools, not active production "
    "runtime dependencies."
)

add_heading(doc, "Data Flow", level=1)
add_monospace_block(doc, """INPUT SOURCES:
- GPS trackers / walk files -> serve.py
- Google Drive -> serve.py or GAS drive_watcher.js
- Google Sheets forecast -> build_weather.py
- Dashboard UI -> serve.py schedule APIs
- Route KML files -> build_dashboard.py

PROCESSING:
- serve.py: serves APIs, validates schedule writes, restores/persists state
- build_weather.py: builds weather advisory JSON
- build_dashboard.py: generates static dashboard HTML

OUTPUT/STORAGE:
- dashboard.html -> team monitoring UI
- schedule_output.json -> claimed schedule assignments
- weather.json -> advisory data
- Walks_Log.txt -> master walk completion log
- GCS bucket -> durable runtime state""", size=9)

add_heading(doc, "Key Data Sources", level=1)
add_table(doc, ["File / Source", "Purpose"], [
    ("Walks_Log.txt", "Master log of completed walks"),
    ("schedule_output.json", "Self-scheduled assignments"),
    ("weather.json", "Weather advisory data"),
    ("Route KMLs", "Geographic route data"),
    ("Availability.xlsx", "Dashboard availability heatmap input"),
    ("Google Drive", "Walk data uploads and sync source"),
])

add_heading(doc, "Technology Stack", level=1)
add_table(doc, ["Layer", "Technology"], [
    ("Backend", "Python 3 HTTP server and pipeline scripts"),
    ("Frontend", "HTML/CSS/JavaScript with Leaflet.js and Chart.js"),
    ("Scheduling", "Self-scheduling APIs backed by schedule_output.json"),
    ("Data Sources", "Google Drive API, Google Sheets, Excel, KML"),
    ("Hosting", "Google Cloud Run with GCS-backed state"),
    ("Dependencies", "google-api-client, openpyxl, pandas, python-docx"),
])

add_heading(doc, "Key Interconnections", level=1)
add_table(doc, ["Component A", "Component B", "Connection"], [
    ("serve.py", "schedule_output.json", "Reads and writes self-scheduled assignments"),
    ("serve.py", "build_dashboard.py", "Triggers dashboard rebuilds"),
    ("build_weather.py", "Google Sheets", "Reads forecast advisory data"),
    ("build_dashboard.py", "schedule_output.json", "Embeds claimed assignments"),
    ("serve.py", "GCS", "Restores and persists runtime state"),
    ("Google Drive", "serve.py", "Drive polling or GAS-triggered sync"),
])

add_heading(doc, "Quick Reference: Execution Flow", level=1)
for step, desc in [
    ("Startup", "serve.py restores GCS state and starts the HTTP server"),
    ("Schedule Claim", "Collector claims a slot in the dashboard; server validates conflicts"),
    ("Forecast Refresh", "GAS calls /api/force-rebuild; server runs weather and dashboard rebuilds"),
    ("Walk Sync", "Drive polling or GAS drive watcher updates walk-log state"),
    ("Dashboard Generation", "build_dashboard.py regenerates dashboard.html"),
    ("Storage", "Durable runtime state is restored from and uploaded to GCS"),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(step + ": ").bold = True
    p.add_run(desc)

add_heading(doc, "For New Team Members", level=1)
add_bullets(doc, [
    "Start with docs/operations/context/CURRENT_STATE.md for current runtime truth",
    "Use docs/operations/context/CLEANUP_PRIORITIES.md for next cleanup tasks",
    "Treat pipelines/_retired/ as historical unless explicitly asked to restore old workflows",
    "Use README.md for setup and API reference",
])

doc.save("Architecture_Overview.docx")
print("Document created successfully: Architecture_Overview.docx")
